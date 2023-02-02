from distutils.util import execute
from docx.enum.text import WD_BREAK
import docx
import sys, os

input= sys.argv[1]
output= sys.argv[2]

# startnum =16
document = docx.Document(f"Document/{input}")
all_paras = document.paragraphs
_number_indecator = 15
def verify_indecator(i):
    global _number_indecator
    #check the question number if it is incremental
    if _number_indecator == i-1:
        _number_indecator =i
        print(_number_indecator)

        return True
    #check if the number have restarted
    elif i==1:
        _number_indecator =i
        return True
    else:
        return False 
def screenq(p):
    #check if the first character of the paragraphs is a number
    if p[:1].isnumeric():
        print(p)
        if p[1:2]=="." or p[1:2]==" ":
            if not p[2:3].isnumeric():
                try:
                    indecator = int(p[:1])
                except ValueError:
                    print(p[:1])
                    indecator=0
                return (verify_indecator(indecator), p[:2])
        elif p[2:3]=="." or p[2:3]==" ":
            if not p[3:4].isnumeric():
                try:
                    indecator = int(p[:2])
                except ValueError:
                    indecator=0
                return (verify_indecator(indecator), p[:3])
        elif p[3:4]=="." or p[3:4]==" ":
            if not p[4:5].isnumeric():
                try:
                    indecator = int(p[:3])
                except ValueError:
                    indecator=0
                return (verify_indecator(indecator), p[:4])
        return (False,0)
    else :
        return (False,0)

    # elif p[:1].isnumeric():
    #     # idn = p[:1]
    #     # inline = p.runs
    #     # for i in range(len(inline)):
    #     #     if idn in inline[i].text:
    #     #         text = inline[i].text.replace(idn, '0'+idn)
    #     #         inline[i].text = text
    #     return True
    
#loop through the paragraph
for idx,para in enumerate(all_paras):
    (is_question ,question_number) = screenq(para.text)
    print(is_question)
    if is_question:
        inline = para.runs
        for i in range(len(inline)):
            if question_number in inline[i].text:
                text = inline[i].text.replace(question_number, '\n')
                inline[i].text = text
        run = all_paras[idx-1].add_run()
        run.add_break(WD_BREAK.PAGE)


document.save("temp.docx")
if not os.path.exists(f"Output/{output}"):
    os.mkdir(f"Output/{output}")
if not os.path.exists(f"Output/{output}-temp"):
    os.mkdir(f"Output/{output}-temp")
os.system(f"python topdf.py temp.docx {output}")
os.remove("temp.docx")

